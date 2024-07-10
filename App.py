import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
from docx import Document
from docx.shared import Inches

# Streamlit app title
st.write('# Automated: DATA ANALYST')

# Uploading excel data and converting into dataframe
uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx")
if uploaded_file is not None:
    # Read the Excel file
    df = pd.read_excel(uploaded_file)
    columns = df.columns
    
    # Display the DataFrame
    st.write("DataFrame:")
    hea = df.head()
    st.write(hea)

    # Create a Word document
    doc = Document()
    doc.add_heading('DataFrame Visualizations', level=1)
    
    # Set the style of seaborn
    sns.set(style="whitegrid")

    # Loop through each column and create different types of plots
    for column in columns:
        # Create a figure for each plot type
        fig, ax = plt.subplots(figsize=(10, 6))
        
        # Histogram
        sns.histplot(df[column], kde=True, ax=ax)
        ax.set_title(f'Histogram of {column}')

        # Save the plot to a BytesIO object
        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        img_stream.seek(0)
        plt.close(fig)

        # Add plot to the Word document
        doc.add_heading(f'Histogram of {column}', level=2)
        doc.add_picture(img_stream, width=Inches(6))
        img_stream.close()

        # Create Boxplot
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.boxplot(data=df[column], ax=ax)
        ax.set_title(f'Boxplot of {column}')

        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        img_stream.seek(0)
        plt.close(fig)

        doc.add_heading(f'Boxplot of {column}', level=2)
        doc.add_picture(img_stream, width=Inches(6))
        img_stream.close()

        # Create Line plot
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.lineplot(data=df[column], ax=ax)
        ax.set_title(f'Line Plot of {column}')

        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        img_stream.seek(0)
        plt.close(fig)

        doc.add_heading(f'Line Plot of {column}', level=2)
        doc.add_picture(img_stream, width=Inches(6))
        img_stream.close()

        # Create Scatter plot
        if len(columns) > 1:
            other_column = columns[(columns.get_loc(column) + 1) % len(columns)]
            fig, ax = plt.subplots(figsize=(10, 6))
            sns.scatterplot(x=df[column], y=df[other_column], ax=ax)
            ax.set_title(f'Scatter Plot of {column} vs {other_column}')

            img_stream = BytesIO()
            plt.savefig(img_stream, format='png')
            img_stream.seek(0)
            plt.close(fig)

            doc.add_heading(f'Scatter Plot of {column} vs {other_column}', level=2)
            doc.add_picture(img_stream, width=Inches(6))
            img_stream.close()

    # Save the Word document to a BytesIO object
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    # Download button for the Word document
    st.download_button(
        label="Download Word Document",
        data=doc_stream,
        file_name="visualizations.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
        